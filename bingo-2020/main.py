from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

import numpy as np
import random

from docx.shared import Pt, Cm

number_range = 75
board_size = 4
magic_number = 7

boards_amount = 2
pages = 1


def generate_board():
    board = [magic_number]
    numbers = range(1, number_range + 1)

    for i in range(board_size ** 2 - 1):
        chosen_number = random.choice(numbers)
        while chosen_number in board:
            chosen_number = random.choice(numbers)
        board.append(chosen_number)

    random.shuffle(board)
    return board


def write_board(document, board):
    table = document.add_table(rows=board_size, cols=board_size)
    table.style = 'Table Grid'

    i = 0
    for row in table.rows:
        for cell in row.cells:
            p = cell.add_paragraph(str(board[i]))
            font = p.runs[0].font

            p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            font.size = Pt(35)

            cell.width = Cm(3)
            row.height = (int) (0.8*cell.width)
            i += 1


document = Document()

for i in range(pages):
    table = document.add_table(rows=boards_amount, cols=boards_amount)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            write_board(cell, generate_board())


document.save('bingo.docx')
